from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Users\Lenovo\Desktop\战棋0.5.2")
OUT = ROOT / "docs" / "开发日志" / "Silent_Reckoning_1987_开发日志_2026-08-24.docx"
QA_IMAGE = ROOT / "art" / "qa" / "unit_animation_four_direction_qa.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "25364A"
MUTED = "667085"
LIGHT = "E8EEF5"
GRID = "CCD5DF"
WHITE = "FFFFFF"


def set_run_font(run, size=11, bold=None, color="000000", italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=GRID, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_bottom_border(paragraph, color=BLUE, size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(ids, default=0) + 1
    nums = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(nums, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "bullet"); lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "•"); lvl.append(lvl_text)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540")
    tabs.append(tab); p_pr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270")
    p_pr.append(ind); lvl.append(p_pr)
    abstract.append(lvl); numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId"); abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id); numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    n_id = OxmlElement("w:numId"); n_id.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, n_id]); p_pr.append(num_pr)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix); set_run_font(r1, bold=True, color=INK)
        r2 = p.add_run(text[len(bold_prefix):]); set_run_font(r2)
    else:
        r = p.add_run(text); set_run_font(r)
    return p


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# Running header/footer.
header_p = section.header.paragraphs[0]
header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
header_p.paragraph_format.space_after = Pt(0)
r = header_p.add_run("SILENT RECKONING · 1987  |  开发日志")
set_run_font(r, size=9, bold=True, color=MUTED)
footer_p = section.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = footer_p.add_run("2026-08-24  |  战棋0.5.2")
set_run_font(r, size=8.5, color=MUTED)

# Memo-style title block.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("昨日更新日志")
set_run_font(r, size=24, bold=True, color=INK)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("《Silent Reckoning · 1987》单位像素动画重制")
set_run_font(r, size=13.5, color=MUTED)

meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(meta, [2700, 6660])
set_table_borders(meta, color=GRID, size="4")
for row, (label, value) in enumerate((
    ("开发日期", "2026年8月24日"),
    ("当前项目", "战棋0.5.2"),
    ("本次主题", "以 T-72B 为唯一标准，统一重制全单位移动与攻击序列帧"),
)):
    set_cell_shading(meta.cell(row, 0), LIGHT)
    p1 = meta.cell(row, 0).paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run(label); set_run_font(r1, size=10.5, bold=True, color=DARK_BLUE)
    p2 = meta.cell(row, 1).paragraphs[0]; p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(value); set_run_font(r2, size=10.5)

summary = doc.add_paragraph()
summary.paragraph_format.space_before = Pt(12)
summary.paragraph_format.space_after = Pt(10)
summary.paragraph_format.line_spacing = 1.25
add_bottom_border(summary, color=BLUE, size="10")
r = summary.add_run("结果摘要：")
set_run_font(r, size=11, bold=True, color=DARK_BLUE)
r = summary.add_run("保留已验收的 T-72B，独立重绘其余24种单位；完成50张移动/攻击图、800个逻辑帧，并通过 Godot 运行测试。")
set_run_font(r, size=11, color=INK)

num_id = add_bullet_numbering(doc)

doc.add_paragraph("一、昨日完成内容", style="Heading 1")
items = [
    ("统一美术基准：", "统一美术基准：将 T-72B 的像素颗粒、近黑轮廓、低色阶、俯视角度、单位占格比例和接地点作为唯一制作标准。"),
    ("全单位重绘：", "全单位重绘：保留 T-72B 原图，其余24种单位全部单独生成母版，不再使用多人合集裁切或由其他兵种换色派生。"),
    ("兵种辨识强化：", "兵种辨识强化：通过炮塔、发射架、火炮、旋翼、无线电天线、探雷器、望远镜等轮廓锚点区分单位；华约和北约采用不同识别色。"),
    ("动画统一：", "动画统一：每种单位均制作移动与攻击序列，固定为四方向×四帧，方向顺序为东南、西南、西北、东北。"),
    ("尺寸与锚点：", "尺寸与锚点：所有游戏成品均为128×128 PNG；单帧32×32；地面接触点统一为单帧坐标(16,29)。"),
    ("问题修复：", "问题修复：首轮检查发现工兵组母版为空透明图，已单独重新生成并通过复查。"),
]
for prefix, text in items:
    add_bullet(doc, num_id, text, prefix)

doc.add_paragraph("二、工具与工程调整", style="Heading 1")
for prefix, text in (
    ("新增重建工具：", "新增重建工具：tools/rebuild_all_units_t72_quality.ps1，可批量重建全部单位，也可通过 -UnitId 单独重建指定单位。"),
    ("保留源文件：", "保留源文件：24种独立母版归档至 art/source/unit_animation_individual_v3，方便后续继续调整。"),
    ("补充制作说明：", "补充制作说明：新增 README_制作规格.md，记录固定尺寸、方向顺序、锚点、提示词和重建流程。"),
    ("运行目录不变：", "运行目录不变：游戏继续从 assets/units/animated/all_units_v2 读取动画，不需要修改现有单位配置路径。"),
):
    add_bullet(doc, num_id, text, prefix)

doc.add_paragraph("三、验证结果", style="Heading 1")
verify = doc.add_table(rows=1, cols=3)
verify.alignment = WD_TABLE_ALIGNMENT.LEFT
set_table_geometry(verify, [2600, 4600, 2160])
set_table_borders(verify)
headers = ("检查项目", "检查结果", "状态")
mark_table_header(verify.rows[0])
for i, text in enumerate(headers):
    set_cell_shading(verify.cell(0, i), LIGHT)
    p = verify.cell(0, i).paragraphs[0]; p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text); set_run_font(r, size=10.2, bold=True, color=DARK_BLUE)
rows = [
    ("图片规格", "50张序列帧图片均为128×128", "通过"),
    ("帧完整性", "800个逻辑帧均包含有效像素", "通过"),
    ("T-72B保护", "移动图文件哈希与基准图一致", "通过"),
    ("Godot导入", "Godot 4.7.2 无头导入退出码为0", "通过"),
    ("项目测试", "Smoke Test：PASS（79 checks）", "通过"),
]
for row_data in rows:
    cells = verify.add_row().cells
    for i, text in enumerate(row_data):
        p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text); set_run_font(r, size=10)
        if i == 2:
            r.font.bold = True; r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
set_table_geometry(verify, [2600, 4600, 2160])

doc.add_paragraph("四、后续建议", style="Heading 1")
for text in (
    "在实际关卡中逐个检查单位与建筑、地块的相对尺寸和遮挡关系。",
    "按武器类型继续细化攻击动作与炮口特效的方向、位置和播放时长。",
    "后续修改优先替换对应单位母版，再使用单单位参数重建，避免影响已经验收的其他单位。",
):
    add_bullet(doc, num_id, text)

if QA_IMAGE.exists():
    doc.add_page_break()
    doc.add_paragraph("附录：全单位四方向效果总览", style="Heading 1")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    shape = p.add_run().add_picture(str(QA_IMAGE), width=Inches(6.35))
    shape._inline.docPr.set("descr", "2026年8月24日全单位四方向像素辨识度检查表，展示25种单位的东南、西南、西北、东北四个方向。")
    shape._inline.docPr.set("title", "全单位四方向效果总览")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(0)
    r = cap.add_run("图1  2026-08-24 全单位四方向像素辨识度检查表")
    set_run_font(r, size=9, italic=True, color=MUTED)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Silent Reckoning 1987 开发日志 2026-08-24"
doc.core_properties.subject = "单位像素动画重制与验证"
doc.core_properties.author = "Silent Reckoning 1987 项目组"
doc.core_properties.keywords = "开发日志, Godot, 像素动画, T-72B"
doc.save(OUT)
print(OUT)
